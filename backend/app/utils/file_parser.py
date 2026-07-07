"""文件解析工具 - PDF/Word/文本解析"""
import os
import logging
from typing import Optional

logger = logging.getLogger(__name__)


def parse_pdf(file_path: str) -> str:
    """解析PDF文件为文本"""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    except Exception as e:
        logger.error(f"PDF解析失败: {e}")
        raise


def parse_docx(file_path: str) -> str:
    """解析Word文件为文本"""
    try:
        from docx import Document
        doc = Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return text.strip()
    except Exception as e:
        logger.error(f"Word解析失败: {e}")
        raise


def parse_text(file_path: str) -> str:
    """解析纯文本文件"""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read().strip()
    except UnicodeDecodeError:
        with open(file_path, "r", encoding="gbk") as f:
            return f.read().strip()


def parse_file(file_path: str) -> str:
    """自动识别文件类型并解析"""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()
    parsers = {
        ".pdf": parse_pdf,
        ".doc": parse_docx,
        ".docx": parse_docx,
        ".txt": parse_text,
    }

    parser = parsers.get(ext)
    if not parser:
        raise ValueError(f"不支持的文件格式: {ext}，支持 PDF/Word/TXT")

    return parser(file_path)
